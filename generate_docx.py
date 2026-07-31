from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Heading styles
for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_style.font.bold = True

doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)

# -- Title Page --
doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Digital Twin Framework for Intelligent Smart Farming\nusing IoT and Cloud Analytics')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('BITE412L - Cloud Computing\nProject Phase-I Report')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph('')

instructor = doc.add_paragraph()
instructor.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = instructor.add_run('Course Instructor: Dr. Priya V')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

submitted = doc.add_paragraph()
submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = submitted.add_run('Submitted by')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

student1 = doc.add_paragraph()
student1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = student1.add_run('Aditya Khanna\t23BIT0324')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

student2 = doc.add_paragraph()
student2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = student2.add_run('Parth Shinde\t\t23BIT0251')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# -- Page break --
doc.add_page_break()

# -- Abstract Section --
doc.add_heading('Abstract', level=1)

abstract_paragraphs = [
    'Modern agriculture faces critical challenges in achieving optimal crop productivity and sustainable resource utilization. Traditional farming practices rely heavily on manual observation and intuition-based decision-making, leading to inefficient water usage, delayed pest and disease detection, and suboptimal crop yield. While existing IoT-based smart farming solutions have introduced real-time sensor monitoring, they remain largely limited to passive data visualization without predictive intelligence or simulation capabilities. Current systems treat environmental factors such as soil health, weather patterns, and crop state as isolated variables rather than dynamically interacting elements, resulting in fragmented and reactive farm management.',

    'To address these limitations, this project proposes a Digital Twin Framework for Intelligent Smart Farming that creates a virtual replica of an agricultural environment, continuously synchronized with real-time IoT sensor data. The framework integrates a simulation engine that enables farmers to conduct "what-if" scenario analysis, such as predicting the impact of drought conditions or delayed irrigation on crop health and yield. An ensemble machine learning model, augmented with derived digital twin state features, provides accurate crop prediction and actionable recommendations. The system implements a closed-loop automation mechanism where predictive insights automatically trigger farming actions such as irrigation adjustments, with the twin state updating to reflect each intervention.',

    'The proposed framework is built entirely on Microsoft Azure cloud services, leveraging Azure IoT Hub for sensor data ingestion, Azure Digital Twins for virtual farm modeling using DTDL twin graphs, Azure Stream Analytics for real-time anomaly detection, Azure Machine Learning for crop prediction model training and deployment, Azure Functions for serverless event-driven automation, Azure Cosmos DB for time-series sensor storage, Azure App Service for hosting the web-based dashboard, Microsoft Entra ID for role-based authentication, and Azure Notification Hubs for delivering smart alerts to farmers.'
]

for para_text in abstract_paragraphs:
    p = doc.add_paragraph(para_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)

# -- Page break --
doc.add_page_break()

# ============================================================
# LITERATURE SURVEY SECTION (Papers 1–15)
# ============================================================
doc.add_heading('Literature Survey', level=1)

# -- Helper to set cell text with font --
def set_cell(cell, text, bold=False, size=9):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

# -- Literature Survey Table --
headers = ['Paper', 'Method', 'Dataset', 'Advantages', 'Limitations', 'Research Gap']

papers_data = [
    [
        'Digital Twins in Agriculture and Forestry: A Review',
        'Systematic literature review of DT applications in agriculture and forestry; categorizes existing architectures and IoT integration approaches.',
        'No primary dataset; reviews and synthesizes findings from multiple published studies.',
        'Comprehensive coverage of DT architectures; identifies key enabling technologies (IoT, cloud, AI); provides taxonomy of agricultural DT use cases.',
        'Review-only with no implementation or experimental validation; does not propose a concrete framework; limited focus on real-time simulation capabilities.',
        'Lacks a unified end-to-end framework that combines DT simulation with ML-based prediction and closed-loop automation for holistic farm management.'
    ],
    [
        'Digital Twin Deployment for Smart Agriculture in Cloud-Fog-Edge Infrastructure',
        'Multi-Agent Systems (MAS) combined with Cloud-Fog-Edge computing architecture for DT deployment in smart farming.',
        'Simulated agricultural IoT data streams across cloud-fog-edge tiers.',
        'Novel distributed architecture handles massive IoT data with low latency; supports real-time DT synchronization; scalable across computing tiers.',
        'Complex multi-tier architecture increases deployment difficulty; no ML-based crop prediction integrated; focuses on infrastructure rather than agricultural intelligence.',
        'Does not integrate machine learning for predictive analytics; missing what-if simulation capabilities for farmer decision support.'
    ],
    [
        'Application Scenarios of Digital Twins for Smart Crop Farming through Cloud-Fog-Edge Infrastructure',
        'Identifies and categorizes DT application scenarios (monitoring, simulation, prediction) within cloud-fog-edge infrastructure for crop farming.',
        'Scenario-based analysis using conceptual crop farming use cases; no specific experimental dataset.',
        'Clearly maps DT application patterns to practical farming use cases; identifies monitoring, simulation, and predictive categories; strong conceptual framework.',
        'Conceptual and scenario-based only — no working prototype or implementation; no quantitative performance evaluation or accuracy metrics.',
        'Application scenarios are identified but not implemented; no ensemble ML integration or automated closed-loop response mechanism demonstrated.'
    ],
    [
        'IoT-Digital Twin-Inspired Smart Irrigation Approach for Optimal Water Utilization',
        'IoT-integrated Digital Twin approach for real-time irrigation optimization; sensor data creates a digital replica of the irrigation system.',
        'Real-time IoT sensor data (soil moisture, temperature, humidity) from experimental irrigation setup.',
        'Practical DT implementation for water management; real-time sensor synchronization; demonstrates measurable water usage reduction.',
        'Focuses solely on irrigation — does not address crop health, yield prediction, or multi-factor farm modeling; single-purpose DT.',
        'Narrow scope limited to water management; does not model the farm as a holistic entity coupling soil, weather, crop state, and yield together.'
    ],
    [
        'Advancing Precision Agriculture Through Digital Twins and Smart Farming Technologies: A Review',
        'Systematic review of 167 studies (2018–2025); proposes a framework for designing and optimizing DTs in smart farming.',
        'No primary dataset; meta-analysis of 167 published studies across multiple databases.',
        'Most comprehensive recent review (167 papers); identifies key adoption barriers; proposes step-by-step DT deployment methodology.',
        'Review paper with no implementation; identified barriers (cost, data integration, interoperability) remain unresolved; no experimental validation.',
        'Highlights the need for standardized, cost-effective DT architectures but does not provide one; no integration of Azure-native cloud services demonstrated.'
    ],
    [
        'Harnessing Digital Twins for Sustainable Agricultural Water Management',
        'Systematic review comparing monitoring DTs vs. predictive DTs; analyzes hybrid modeling (physics-based + data-driven ML) for water management.',
        'No primary dataset; reviews existing studies on DT-based agricultural water management systems.',
        'Clearly distinguishes monitoring vs. predictive DT approaches; highlights hybrid modeling as superior; identifies AI/ML as key enabler.',
        'Review-only; does not implement a predictive DT; limited to water management domain without broader farm system integration.',
        'Identifies predictive DTs as the future direction but no working implementation exists that couples predictive DT with ensemble ML and automated farming actions.'
    ],
    [
        'Enhancing Crop Yield Predictions with PEnsemble4: IoT and ML-Driven for Precision Agriculture',
        'PEnsemble4 — weighted ensemble ML model (combining multiple classifiers) integrated with IoT sensor and UAV imagery data for maize yield prediction.',
        'UAV-captured vegetation imagery (CIre, NDRE indices) combined with IoT environmental sensor data (soil, nutrients, weather) for maize fields.',
        'Ensemble approach outperforms individual models; achieves 91% accuracy; enables earlier yield prediction (R2 stage vs conventional R6 stage).',
        'Tested only on maize — generalizability to other crops unverified; no digital twin integration; standalone ML pipeline without real-time farm synchronization.',
        'Ensemble ML is proven effective but not integrated into a digital twin framework; lacks real-time twin state features as additional model inputs.'
    ],
    [
        'IoT-Enabled Soil Nutrient Analysis and Crop Recommendation System',
        'End-to-end IoT pipeline for real-time soil nutrient classification (NPK, pH, temperature, moisture) with Random Forest ML for crop recommendation.',
        'Standard crop recommendation dataset (Kaggle) with features: N, P, K, pH, temperature, humidity, rainfall.',
        'Real-time soil analysis via IoT sensors; high classification accuracy with Random Forest; practical and deployable system.',
        'Standalone classification system — no digital twin, no simulation, no cloud-native architecture; uses basic single-model ML rather than ensemble.',
        'Treats crop recommendation as an isolated classification task; does not integrate soil analysis into a broader digital twin with simulation and closed-loop automation.'
    ],
    [
        'AgriTwin-Sim: An Interactive Digital Twin Framework for AI-Driven Smart Farming',
        'Interactive web-based DT framework integrating a high-fidelity simulation core with an AI engine for diagnostics, prognostics, and prescriptive analytics in a greenhouse setting.',
        'Virtual sensor data generated by a simulated IoT network emulating greenhouse environmental parameters (temperature, humidity, CO2, soil moisture, light).',
        'Complete virtual-sensor-to-intelligent-action pipeline; supports what-if scenario exploration; web-based interface makes complex DT simulations accessible and interactive.',
        'Relies entirely on simulated/synthetic data — not validated with real-world field deployments; greenhouse-specific model may not generalize to open-field farming.',
        'Framework architecture supports RL-based control but does not implement or evaluate ensemble ML for crop yield prediction integrated with twin state features.'
    ],
    [
        'Digital Twins in Agriculture: Orchestration and Applications',
        'Critical review of DT orchestration across the agricultural lifecycle (edaphic, phytotechnologic, postharvest, farm infrastructure) with supporting case studies.',
        'No primary dataset; review synthesizes case studies and published DT deployments across agricultural production and supply chain domains.',
        'Comprehensive lifecycle coverage from soil to supply chain; identifies direct benefits of DTs for resource optimization (water, fertilizer, pesticide savings).',
        'Review-focused with no novel implementation; highlights data management complexity but does not propose a concrete solution architecture.',
        'Lack of standardized orchestration protocols across cloud platforms; no integration of ensemble ML or closed-loop automation demonstrated.'
    ],
    [
        'Conceptual Framework for Smart 4.0 Hydroponic Farming: A Data-Driven Approach to Sustainable Agriculture',
        'Industry 4.0-driven conceptual framework integrating IoT sensors, ML, and Big Data Analytics for fully automated hydroponic farming with real-time monitoring.',
        'Hydroponic testbed data including pH, EC (electrical conductivity), nutrient levels, and ambient environmental conditions.',
        'Achieves 3x crop yield per unit area vs conventional methods; automated dashboards and AI-powered decision-making improve operational efficiency.',
        'Specific to hydroponic/vertical farming — not directly applicable to traditional soil-based agriculture; conceptual framework with limited large-scale commercial validation.',
        'Integration of digital twin simulation with ensemble ML for predictive yield optimization in hydroponics remains unexplored.'
    ],
    [
        'AI-Powered Predictive Maintenance System for Smart Agriculture Using IoT and Cloud Computing',
        'Cloud-native predictive maintenance system using Azure IoT Hub and Azure ML for real-time equipment health monitoring, time-series forecasting, and anomaly detection.',
        'Real-time telemetry data from agricultural machinery (tractors, irrigation pumps) — vibration, temperature, and operational hour logs collected via IoT sensors.',
        'Azure-native architecture (IoT Hub + ML); reduces unplanned equipment downtime; scalable and applicable to various farm machinery categories.',
        'Focuses exclusively on machinery maintenance — does not address biological crop modeling, yield prediction, or environmental simulation.',
        'Does not link equipment health insights to overall crop yield outcomes within a unified digital twin framework.'
    ],
    [
        'A Cloud-IoT Framework for Smart Farming and Sustainable Agri-Food Management',
        'Adaptive Cloud-IoT framework combining cloud-based federated intelligence with edge-layer pre-processing for real-time crop monitoring and yield optimization.',
        'Multi-sensor IoT deployment data spanning soil moisture, weather, and crop health parameters across experimental farm sites.',
        '32% improvement in data transmission efficiency; 21% reduction in cloud processing latency; 19% better yield estimation accuracy than traditional approaches.',
        'Broad farm-to-fork scope limits depth in predictive crop modeling; complex multi-stakeholder deployment across heterogeneous environments.',
        'Missing a closed-loop digital twin simulation engine for the farming phase; no what-if scenario analysis capability.'
    ],
    [
        'A Smart Agricultural Model by Integrating IoT, Mobile and Cloud-based Big Data Analytics',
        'End-to-end smart agriculture model integrating IoT sensing, cloud-based Big Data analytics (Hadoop/MapReduce), and mobile app delivery for farmer decision support.',
        'IoT sensor data (soil properties, environmental parameters) stored in cloud; analyzed using Big Data mining techniques for crop and market prediction.',
        'Accessible via mobile app for farmers; integrates IoT data collection with cloud-based big data analytics; covers crop prediction and market analysis.',
        'Relies on passive monitoring and analytics rather than interactive simulations; lacks automated actuation or closed-loop control mechanisms.',
        'No digital twin component; absence of predictive twin simulation capable of what-if scenario testing for real-time farming decisions.'
    ],
    [
        'A Multi-Stakeholder Collaborative Approach for Smart Agriculture Leveraging IoT, Data Analytics, and Precision Farming Techniques',
        'Multi-layered IoT-Cloud-Dew architecture using LEACH algorithm for energy-efficient sensor data routing combined with drone-based aerial monitoring for precision farming.',
        'Simulated sensor network (20 nodes, 5 rounds) measuring soil moisture, temperature, humidity, and crop health; drone-captured aerial imagery.',
        'Energy-efficient LEACH-based clustering extends sensor network lifetime; three-layer architecture (IoT-Dew-Cloud) balances latency and computational cost; drone integration enhances monitoring coverage.',
        'Small-scale simulation (20 nodes); Dew computing layer adds architectural complexity; no ML-based crop prediction or yield forecasting integrated.',
        'Does not incorporate digital twin simulation or ensemble ML; lacks closed-loop automation where predictive insights trigger farming actions automatically.'
    ]
]

# Create table
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'

# Set header row
for i, header in enumerate(headers):
    set_cell(table.rows[0].cells[i], header, bold=True, size=10)

# Add data rows
for row_data in papers_data:
    row = table.add_row()
    for i, cell_text in enumerate(row_data):
        set_cell(row.cells[i], cell_text, bold=False, size=9)

# Set column widths
from docx.shared import Cm
col_widths = [Cm(3.5), Cm(3.5), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0)]
for row in table.rows:
    for i, width in enumerate(col_widths):
        row.cells[i].width = width

# -- Page break --
doc.add_page_break()

# ============================================================
# PROJECT OBJECTIVES SECTION
# ============================================================
doc.add_heading('Project Objectives', level=1)

objectives = [
    'To design and develop a Digital Twin framework that creates a virtual replica of an agricultural environment, continuously synchronized with real-time IoT sensor data ingested through Azure IoT Hub, enabling holistic farm state monitoring.',

    'To implement a "what-if" simulation engine within the Digital Twin that allows farmers to conduct scenario analysis — such as predicting the impact of drought conditions, delayed irrigation, or nutrient deficiency — on crop health and yield before taking action.',

    'To build and deploy an ensemble machine learning model using Azure Machine Learning that integrates derived digital twin state features with environmental sensor data, achieving improved crop yield prediction accuracy over single-model approaches.',

    'To establish a closed-loop automation mechanism using Azure Functions where predictive insights from the ML model automatically trigger farming actions (e.g., irrigation adjustments, fertilizer alerts), with the digital twin state updating in real time to reflect each intervention.',

    'To architect the entire system as a cloud-native solution on Microsoft Azure, leveraging Azure Digital Twins for DTDL-based twin graph modeling, Azure Stream Analytics for real-time anomaly detection, Azure Cosmos DB for time-series storage, and Azure Notification Hubs for delivering smart alerts to farmers.',

    'To develop an interactive web-based dashboard hosted on Azure App Service, secured with Microsoft Entra ID role-based authentication, providing farmers with real-time visualization of farm conditions, prediction charts, simulation controls, and alert management capabilities.'
]

for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Bold number
    run_num = p.add_run(f'{i}. ')
    run_num.font.name = 'Times New Roman'
    run_num.font.size = Pt(12)
    run_num.font.bold = True
    # Objective text
    run_text = p.add_run(obj)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)

# -- Page break --
doc.add_page_break()

# ============================================================
# NOVELTY SUMMARY SECTION
# ============================================================
doc.add_heading('Novelty Summary', level=1)

novelty_intro = (
    'The literature survey reveals that existing smart farming solutions fall into '
    'two broad categories: IoT-based monitoring systems that passively display sensor data '
    'without predictive intelligence, and digital twin frameworks that remain largely '
    'conceptual or limited to single-purpose applications such as irrigation management. '
    'No existing work presents an end-to-end, cloud-native digital twin framework that '
    'unifies real-time simulation, ensemble machine learning, and closed-loop automation '
    'for holistic farm management. The proposed framework addresses this gap through the '
    'following novel contributions:'
)
p = doc.add_paragraph(novelty_intro)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Inches(0.5)

novelty_points = [
    (
        'What-If Simulation Engine',
        'Unlike existing monitoring-only systems, the framework integrates a simulation engine '
        'directly within the Azure Digital Twins service, enabling farmers to conduct scenario '
        'analysis — such as predicting the impact of drought conditions, delayed irrigation, '
        'or nutrient deficiency on crop health and yield — before committing to any farming action. '
        'This transforms the digital twin from a passive data mirror into an active decision-support tool.'
    ),
    (
        'Ensemble ML Augmented with Twin State Features',
        'While prior works such as PEnsemble4 have demonstrated the effectiveness of ensemble machine '
        'learning for crop prediction, they operate as standalone ML pipelines without digital twin '
        'integration. This framework uniquely augments the ensemble model with derived digital twin '
        'state features (e.g., cumulative stress indices, simulated growth trajectories), providing '
        'richer context that improves prediction accuracy beyond what raw sensor data alone can achieve.'
    ),
    (
        'Closed-Loop Automation with Twin State Feedback',
        'Existing systems treat prediction and actuation as separate, disconnected processes. This '
        'framework implements a true closed-loop mechanism where predictive insights from the ML model '
        'automatically trigger farming actions (such as irrigation adjustments or fertilizer alerts) '
        'through Azure Functions, and the digital twin state updates in real time to reflect each '
        'intervention — creating a continuous feedback cycle between prediction, action, and twin synchronization.'
    ),
    (
        'End-to-End Azure-Native Architecture',
        'While individual Azure services have been used in prior agricultural IoT projects, no existing '
        'work demonstrates a fully integrated Azure-native stack purpose-built for agricultural digital '
        'twins. The proposed architecture cohesively combines Azure IoT Hub, Azure Stream Analytics, '
        'Azure Digital Twins (DTDL twin graphs), Azure Machine Learning, Azure Functions, Azure Cosmos DB, '
        'Azure App Service, Microsoft Entra ID, and Azure Notification Hubs into a unified, production-ready framework.'
    ),
    (
        'Holistic Farm Modeling',
        'Prior digital twin implementations in agriculture tend to address isolated aspects — irrigation '
        'management, soil monitoring, or equipment maintenance — as independent systems. This framework '
        'models the farm as a holistic entity where soil health, weather patterns, crop state, and yield '
        'interact as dynamically coupled variables within a single twin graph, enabling system-level '
        'insights that fragmented approaches cannot provide.'
    )
]

for title, description in novelty_points:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Bold title
    run_title = p.add_run(f'{title}: ')
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(12)
    run_title.font.bold = True
    # Description
    run_desc = p.add_run(description)
    run_desc.font.name = 'Times New Roman'
    run_desc.font.size = Pt(12)

# -- Page break --
doc.add_page_break()

# ============================================================
# PROPOSED ARCHITECTURE SECTION
# ============================================================
doc.add_heading('Proposed Architecture', level=1)

# Diagram 1 text
doc.add_heading('Diagram 1: Azure Cloud Architecture', level=2)
p_cloud_desc = doc.add_paragraph(
    'The Azure Cloud Architecture diagram illustrates the detailed interactions between the core '
    'Azure services forming the backbone of the Digital Twin framework. The data flow originates '
    'from the IoT sensors (Temperature, Soil Moisture, Humidity, pH) and enters the cloud via '
    'Azure IoT Hub. Real-time processing and anomaly detection are handled by Azure Stream Analytics, '
    'which feeds into Azure Functions for event-driven logic execution. The central component is '
    'Azure Digital Twins, which maintains the twin model and graph. This state is actively utilized '
    'by the What-If Simulation Engine for scenario analysis and by Azure Machine Learning for ensemble '
    'crop prediction. Azure Functions also route data to Azure Cosmos DB (time-series storage) and '
    'Azure Blob Storage (unstructured data). The closed-loop system is completed with automated actions '
    'triggered from the ML module back to Azure Functions, and smart alerts dispatched via Azure '
    'Notification Hub. Security and access control for the web dashboard are managed by Microsoft Entra ID, '
    'with overall diagnostics monitored by Azure Monitor.'
)
p_cloud_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Embed Diagram 1
p_cloud_img = doc.add_paragraph()
p_cloud_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_cloud_img = p_cloud_img.add_run()
run_cloud_img.add_picture(r'c:\Users\MSI\Desktop\Aditya\Cloud Computing Project\cloud_architecture.png', width=Inches(6.0))

doc.add_paragraph()

# Diagram 2 text
doc.add_heading('Diagram 2: Complete System Architecture', level=2)
p_system_desc = doc.add_paragraph(
    'The Complete System Architecture provides an end-to-end view of the project workflow, '
    'structured across three distinct tiers: the IoT Layer, the Cloud Platform, and the Application Layer. '
    'The IoT Layer comprises a Sensor Data Generator that simulates environmental readings (temperature, '
    'humidity, soil moisture, pH) representing the physical farm. This data is securely ingested into the '
    'Cloud Platform via Azure IoT Hub, processed through Stream Analytics and Functions, and used to update '
    'the Azure Digital Twin model. The cloud tier also hosts the storage, machine learning, and simulation '
    'engines required to generate predictive insights and automated actions. Finally, the Application Layer '
    'exposes these capabilities to the farmer through a Web Dashboard hosted on Azure App Service. This '
    'interactive interface allows users to view real-time farm conditions, manage alerts from the Notification '
    'Service, and directly execute "what-if" simulations to inform their farming strategies.'
)
p_system_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Embed Diagram 2
p_system_img = doc.add_paragraph()
p_system_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_system_img = p_system_img.add_run()
run_system_img.add_picture(r'c:\Users\MSI\Desktop\Aditya\Cloud Computing Project\system_architecture.png', width=Inches(6.0))

# -- Page break --
doc.add_page_break()

# ============================================================
# DATASET DETAILS SECTION
# ============================================================
doc.add_heading('Dataset Details', level=1)

p_data_desc = doc.add_paragraph(
    'To train the foundational machine learning model for crop yield prediction and health analysis '
    'within the Azure Machine Learning workspace, the following dataset has been selected. This dataset '
    'serves as the baseline historical data, which will be augmented with live, synthetic sensor streams '
    'during the digital twin simulation phase.'
)
p_data_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Dataset attributes
dataset_info = [
    ('Dataset Name', 'Crop Recommendation Dataset'),
    ('Source', 'Kaggle (compiled from Indian Council of Agricultural Research data)'),
    ('URL', 'https://www.kaggle.com/datasets/atharvaingle/crop-recommendation-dataset'),
    ('Size', '~150 KB (CSV format)'),
    ('Number of Records', '2,200 rows'),
    ('Number of Features', '8 columns (N, P, K, temperature, humidity, pH, rainfall, label)'),
    ('Data Type', 'Tabular, Time-Series adaptable'),
    ('License', 'Database: Open Database, Contents: Database Contents (DbCL)'),
    ('Purpose', 'Provides foundational ML training data to map environmental constraints (soil nutrients, weather) to optimal crop types, which is essential for the predictive engine of the digital twin.'),
    ('Preprocessing Required', 'Missing value imputation (if any), scaling/normalization of continuous environmental variables, label encoding for crop categories, and synthetic time-series augmentation to simulate live IoT streams.')
]

# Create table
table_dataset = doc.add_table(rows=1, cols=2)
table_dataset.style = 'Table Grid'

# Set header row
set_cell(table_dataset.rows[0].cells[0], 'Attribute', bold=True, size=11)
set_cell(table_dataset.rows[0].cells[1], 'Description', bold=True, size=11)

# Add data rows
for attr, desc in dataset_info:
    row = table_dataset.add_row()
    set_cell(row.cells[0], attr, bold=True, size=10)
    set_cell(row.cells[1], desc, bold=False, size=10)

# Set column widths
col_widths_ds = [Cm(5.0), Cm(11.0)]
for row in table_dataset.rows:
    for i, width in enumerate(col_widths_ds):
        row.cells[i].width = width

# -- Page break --
doc.add_page_break()

# ============================================================
# AZURE SERVICES PLANNING SECTION
# ============================================================
doc.add_heading('Azure Services Planning', level=1)

p_services_desc = doc.add_paragraph(
    'The proposed Digital Twin Framework relies on a robust, cloud-native architecture built entirely '
    'on Microsoft Azure. The following table details the specific Azure services utilized and their '
    'exact purpose within the project ecosystem:'
)
p_services_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

azure_services = [
    ('Azure IoT Hub', 'Securely ingests massive volumes of real-time telemetry data from the simulated agricultural sensors and routes it to downstream processing services.'),
    ('Azure Stream Analytics', 'Performs real-time processing and anomaly detection on the incoming IoT data streams before they reach the storage and twin layers.'),
    ('Azure Functions', 'Serves as the event-driven serverless compute engine. It auto-processes incoming data, triggers the digital twin updates, and executes closed-loop automated farming actions.'),
    ('Azure Digital Twins', 'The core component hosting the DTDL (Digital Twins Definition Language) twin models and graphs. It maintains the live virtual replica of the farm and facilitates "what-if" scenario simulations.'),
    ('Azure Machine Learning', 'Trains, evaluates, and deploys the ensemble ML model for crop yield prediction, integrating real-time twin state features for enhanced accuracy.'),
    ('Azure Cosmos DB', 'Provides low-latency, highly scalable NoSQL storage for the continuous influx of time-series IoT sensor data.'),
    ('Azure Blob Storage', 'Stores unstructured data such as historical datasets (e.g., the Kaggle Crop Recommendation dataset), ML model artifacts, and visual assets.'),
    ('Azure App Service', 'Provides the scalable hosting environment for the interactive web-based dashboard used by the farmers.'),
    ('Microsoft Entra ID', 'Ensures secure authentication and role-based access control (RBAC) to the web dashboard and underlying cloud resources.'),
    ('Azure Notification Hubs', 'Dispatches smart alerts and emergency notifications (e.g., critical drought warnings or fertilizer needs) across various platforms (mobile/web) to the farmer.'),
    ('Azure Monitor', 'Continuously monitors the health, performance, and availability of all integrated Azure services and application components.')
]

# Create table
table_services = doc.add_table(rows=1, cols=2)
table_services.style = 'Table Grid'

# Set header row
set_cell(table_services.rows[0].cells[0], 'Azure Service', bold=True, size=11)
set_cell(table_services.rows[0].cells[1], 'Purpose in Project', bold=True, size=11)

# Add data rows
for service, purpose in azure_services:
    row = table_services.add_row()
    set_cell(row.cells[0], service, bold=True, size=10)
    set_cell(row.cells[1], purpose, bold=False, size=10)

# Set column widths
col_widths_svc = [Cm(4.0), Cm(12.0)]
for row in table_services.rows:
    for i, width in enumerate(col_widths_svc):
        row.cells[i].width = width

# -- Save --
output_path = r'c:\Users\MSI\Desktop\Aditya\Cloud Computing Project\Cloud_Computing_Project_Report.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')




